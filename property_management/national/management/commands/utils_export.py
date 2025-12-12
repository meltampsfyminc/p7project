from io import BytesIO
from django.template.loader import render_to_string
from xhtml2pdf import pisa
from docx import Document
from django.http import HttpResponse
import datetime


def render_to_pdf(template_src, context):
    html = render_to_string(template_src, context)
    result = BytesIO()
    pdf = pisa.pisaDocument(BytesIO(html.encode("UTF-8")), result)

    if not pdf.err:
        return result.getvalue()
    return None


def render_to_docx(context):
    document = Document()
    document.add_heading("National Report", level=1)

    document.add_paragraph(
        f"District: {context['district'].name} ({context['district'].dcode})"
    )
    document.add_paragraph(
        f"Local: {context['local'].name} ({context['local'].lcode})"
    )
    document.add_paragraph(f"Year: {context['report'].year}")
    document.add_paragraph(f"Generated: {datetime.datetime.now()}")

    summary = context.get("summary")

    if summary:
        table = document.add_table(rows=1, cols=6)
        hdr = table.rows[0].cells
        hdr[0].text = "P1"
        hdr[1].text = "P2"
        hdr[2].text = "P3"
        hdr[3].text = "P4"
        hdr[4].text = "P5"
        hdr[5].text = "Total"

        row = table.add_row().cells
        row[0].text = str(summary.p1_total or "")
        row[1].text = str(summary.p2_total or "")
        row[2].text = str(summary.p3_total or "")
        row[3].text = str(summary.p4_total or "")
        row[4].text = str(summary.p5_total or "")
        row[5].text = str(summary.total_summary or "")

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
